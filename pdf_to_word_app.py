import tkinter as tk
from tkinter import filedialog, ttk, messagebox
from pdf2docx import Converter
from deep_translator import GoogleTranslator
from docx import Document
import threading
import os
import logging

# Suppress pdf2docx logging noise
logging.getLogger("pdf2docx").setLevel(logging.WARNING)


class PDFtoWordApp:
    def __init__(self, root):
        self.root = root
        self.root.title("PDF to Word (PT)")
        self.root.geometry("520x320")
        self.root.resizable(False, False)
        self.root.configure(bg="#f0f0f0")

        self.pdf_path = tk.StringVar()
        self.translate_var = tk.BooleanVar(value=True)

        # Title
        tk.Label(root, text="Conversor PDF para Word", font=("Segoe UI", 16, "bold"),
                 bg="#f0f0f0", fg="#333").pack(pady=(20, 5))
        tk.Label(root, text="Com tradução para Português (PT)", font=("Segoe UI", 10),
                 bg="#f0f0f0", fg="#666").pack(pady=(0, 15))

        # File selection frame
        frame = tk.Frame(root, bg="#f0f0f0")
        frame.pack(padx=20, fill="x")

        self.file_label = tk.Label(frame, text="Nenhum ficheiro selecionado",
                                   font=("Segoe UI", 9), bg="#fff", fg="#999",
                                   relief="solid", bd=1, padx=10, pady=8, anchor="w")
        self.file_label.pack(side="left", fill="x", expand=True)

        tk.Button(frame, text="Selecionar PDF", font=("Segoe UI", 9, "bold"),
                  bg="#0078d4", fg="white", relief="flat", padx=15, pady=8,
                  cursor="hand2", command=self.select_file).pack(side="right", padx=(10, 0))

        # Translate checkbox
        tk.Checkbutton(root, text="Traduzir para Português (PT)",
                       variable=self.translate_var, font=("Segoe UI", 10),
                       bg="#f0f0f0", activebackground="#f0f0f0").pack(pady=(15, 5))

        # Convert button
        self.convert_btn = tk.Button(root, text="Converter", font=("Segoe UI", 11, "bold"),
                                     bg="#107c10", fg="white", relief="flat",
                                     padx=40, pady=10, cursor="hand2",
                                     command=self.start_conversion)
        self.convert_btn.pack(pady=(10, 10))

        # Progress bar
        self.progress = ttk.Progressbar(root, length=460, mode="determinate")
        self.progress.pack(padx=30, pady=(5, 5))

        # Status label
        self.status = tk.Label(root, text="", font=("Segoe UI", 9), bg="#f0f0f0", fg="#333")
        self.status.pack(pady=(0, 10))

    def select_file(self):
        path = filedialog.askopenfilename(
            title="Selecionar ficheiro PDF",
            filetypes=[("Ficheiros PDF", "*.pdf")]
        )
        if path:
            self.pdf_path.set(path)
            name = os.path.basename(path)
            self.file_label.config(text=name, fg="#333")

    def set_status(self, text):
        self.status.config(text=text)
        self.root.update_idletasks()

    def set_progress(self, value):
        self.progress["value"] = value
        self.root.update_idletasks()

    def start_conversion(self):
        if not self.pdf_path.get():
            messagebox.showwarning("Aviso", "Selecione um ficheiro PDF primeiro.")
            return
        self.convert_btn.config(state="disabled")
        self.set_progress(0)
        threading.Thread(target=self.convert, daemon=True).start()

    def convert(self):
        try:
            pdf = self.pdf_path.get()
            folder = os.path.dirname(pdf)
            name = os.path.splitext(os.path.basename(pdf))[0]

            if self.translate_var.get():
                output = os.path.join(folder, f"{name}_PT.docx")
            else:
                output = os.path.join(folder, f"{name}.docx")

            # Step 1: Convert PDF to Word
            self.set_status("A converter PDF para Word...")
            self.set_progress(10)

            cv = Converter(pdf)
            cv.convert(output)
            cv.close()

            self.set_progress(40)

            # Step 2: Translate if checked
            if self.translate_var.get():
                self.set_status("A traduzir para Português (PT)...")
                doc = Document(output)
                translator = GoogleTranslator(source="en", target="pt")

                # Count total runs to translate for progress
                all_paras = list(doc.paragraphs)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            all_paras.extend(cell.paragraphs)

                total = len(all_paras)
                done = 0

                for para in doc.paragraphs:
                    self._translate_para(para, translator)
                    done += 1
                    self.set_progress(40 + int((done / total) * 50))

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                self._translate_para(para, translator)
                                done += 1
                                self.set_progress(40 + int((done / total) * 50))

                doc.save(output)

            self.set_progress(100)
            self.set_status(f"Concluído: {os.path.basename(output)}")
            self.root.after(0, lambda: messagebox.showinfo(
                "Sucesso", f"Ficheiro guardado em:\n{output}"))

        except Exception as e:
            self.set_status("Erro!")
            self.root.after(0, lambda: messagebox.showerror("Erro", str(e)))
        finally:
            self.root.after(0, lambda: self.convert_btn.config(state="normal"))

    def _translate_para(self, para, translator):
        if not para.text.strip():
            return
        for run in para.runs:
            text = run.text
            if not text or not text.strip():
                continue
            leading = len(text) - len(text.lstrip())
            trailing = len(text) - len(text.rstrip())
            lead = text[:leading] if leading else ""
            trail = text[len(text) - trailing:] if trailing else ""
            core = text.strip()
            if not core:
                continue
            try:
                translated = translator.translate(core)
                if translated:
                    run.text = lead + translated + trail
            except Exception:
                pass


if __name__ == "__main__":
    root = tk.Tk()
    app = PDFtoWordApp(root)
    root.mainloop()
