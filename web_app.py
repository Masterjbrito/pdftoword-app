import io
import logging
import os
import shutil
import subprocess
import tempfile
import uuid
import zipfile
from pathlib import Path

import fitz
from deep_translator import GoogleTranslator
from docx import Document
from flask import Flask, after_this_request, redirect, render_template, request, send_file, url_for
from pdf2docx import Converter
from PIL import Image
from pypdf import PdfReader, PdfWriter
from werkzeug.utils import secure_filename
import img2pdf

try:
    import pytesseract
except Exception:
    pytesseract = None

try:
    import speech_recognition as sr
except Exception:
    sr = None

try:
    import yt_dlp
except Exception:
    yt_dlp = None

try:
    from pydub import AudioSegment
except Exception:
    AudioSegment = None

try:
    from docx2pdf import convert as docx2pdf_convert
except Exception:
    docx2pdf_convert = None

logging.getLogger("pdf2docx").setLevel(logging.WARNING)

BASE_DIR = Path(__file__).resolve().parent
UPLOAD_DIR = BASE_DIR / "uploads"
OUTPUT_DIR = BASE_DIR / "outputs"
TEMP_DIR = BASE_DIR / "temp"
RUNTIME_DIR = Path(tempfile.gettempdir()) / "pdftoword_runtime"

for folder in (UPLOAD_DIR, OUTPUT_DIR, TEMP_DIR, RUNTIME_DIR):
    folder.mkdir(exist_ok=True)

TESSERACT_AVAILABLE = shutil.which("tesseract") is not None and pytesseract is not None

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 40 * 1024 * 1024  # 40MB

TOOLS = [
    ("pdf-to-word", "PDF to Word"),
    ("word-to-pdf", "Word to PDF"),
    ("pdf-to-images", "PDF to Images"),
    ("images-to-pdf", "Images to PDF"),
    ("merge-pdf", "Merge PDFs"),
    ("split-pdf", "Split PDF"),
    ("compress-pdf", "Compress PDF"),
    ("protect-pdf", "Protect PDF"),
    ("unlock-pdf", "Unlock PDF"),
    ("sign-pdf", "Sign PDF"),
    ("remove-watermark", "Remove Watermark"),
    ("remove-password", "Remove Document Password"),
    ("ocr", "OCR / Text Extraction"),
    ("translate", "Translate Document"),
    ("transcribe-audio", "Transcribe Audio to TXT"),
    ("youtube-to-mp3", "YouTube to MP3"),
    ("youtube-to-mp4", "YouTube to MP4"),
    ("youtube-playlist-to-mp3", "YouTube Playlist to MP3"),
    ("youtube-playlist-to-mp4", "YouTube Playlist to MP4"),
    ("spotify-to-mp3", "Spotify to MP3"),
    ("spotify-playlist-to-mp3", "Spotify Playlist to MP3"),
]
TOOL_MAP = {slug: title for slug, title in TOOLS}
TOOL_DESC = {
    "pdf-to-word": "Convert PDF into editable DOCX, with optional translation.",
    "word-to-pdf": "Convert DOCX to PDF.",
    "pdf-to-images": "Export each PDF page as image.",
    "images-to-pdf": "Merge images into a single PDF.",
    "merge-pdf": "Combine multiple PDFs into one file.",
    "split-pdf": "Split PDF by pages or range.",
    "compress-pdf": "Reduce PDF file size.",
    "protect-pdf": "Apply a password to a PDF.",
    "unlock-pdf": "Remove PDF protection with a valid password.",
    "sign-pdf": "Add text and/or image signature into PDF.",
    "remove-watermark": "Try to remove text watermarks from PDF.",
    "remove-password": "Remove PDF password when current password is provided.",
    "ocr": "Extract text from PDFs and images.",
    "translate": "Translate DOCX, PDF or TXT documents.",
    "transcribe-audio": "Transcribe audio/video into TXT.",
    "youtube-to-mp3": "Extract audio from YouTube link as MP3.",
    "youtube-to-mp4": "Download YouTube video as MP4.",
    "spotify-to-mp3": "Convert Spotify links to MP3 (spotdl).",
    "youtube-playlist-to-mp3": "Convert YouTube playlist to MP3 ZIP.",
    "youtube-playlist-to-mp4": "Convert YouTube playlist to MP4 ZIP.",
    "spotify-playlist-to-mp3": "Convert Spotify playlist to MP3 ZIP.",
}
CATEGORY_ITEMS = {
    "conversao": ["pdf-to-word", "word-to-pdf", "pdf-to-images", "images-to-pdf"],
    "edicao": ["merge-pdf", "split-pdf", "compress-pdf", "sign-pdf", "remove-watermark"],
    "seguranca": ["protect-pdf", "unlock-pdf", "remove-password"],
    "texto": ["ocr", "translate", "transcribe-audio"],
    "media": [
        "youtube-to-mp3",
        "youtube-to-mp4",
        "youtube-playlist-to-mp3",
        "youtube-playlist-to-mp4",
        "spotify-to-mp3",
        "spotify-playlist-to-mp3",
    ],
}
FFMPEG_PATH = shutil.which("ffmpeg")
LIBREOFFICE_PATH = shutil.which("libreoffice") or shutil.which("soffice")


@app.context_processor
def inject_globals():
    # Filter out media category for AdSense approval phase
    filtered_category_items = {k: v for k, v in CATEGORY_ITEMS.items() if k != "media"}
    hidden_slugs = CATEGORY_ITEMS.get("media", [])
    filtered_tools = [t for t in TOOLS if t[0] not in hidden_slugs]
    filtered_tool_map = {slug: title for slug, title in filtered_tools}

    return {
        "adsense_id": os.environ.get("ADSENSE_CLIENT_ID", "ca-pub-3114217198436430"),
        "base_url": os.environ.get("BASE_URL", request.host_url.rstrip("/")),
        "category_items": filtered_category_items,
        "tool_map": filtered_tool_map,
        "tools": filtered_tools,
        "tesseract_available": TESSERACT_AVAILABLE,
    }


def unique_name(stem: str, suffix: str) -> str:
    safe_stem = secure_filename(stem) or "file"
    return f"{safe_stem}_{uuid.uuid4().hex[:10]}{suffix}"


def safe_download_name(raw_name: str, ext: str, fallback: str) -> str:
    name = (raw_name or "").strip()
    if not name:
        name = fallback
    invalid = '<>:"/\\|?*'
    for char in invalid:
        name = name.replace(char, " ")
    name = " ".join(name.split()).strip(". ")
    if not name:
        name = fallback
    return f"{name}{ext}"


def zip_files(files: list[Path], zip_path: Path) -> None:
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for item in files:
            zf.write(item, item.name)


def queue_cleanup(paths: list[Path], dirs: list[Path] | None = None):
    dirs = dirs or []

    @after_this_request
    def cleanup(response):
        for path in paths:
            try:
                path.unlink(missing_ok=True)
            except Exception:
                pass
        for directory in dirs:
            try:
                shutil.rmtree(directory, ignore_errors=True)
            except Exception:
                pass
        return response


def translate_docx(docx_path: Path, source_lang: str, target_lang: str) -> None:
    for key in ("http_proxy", "https_proxy", "HTTP_PROXY", "HTTPS_PROXY", "ALL_PROXY", "all_proxy"):
        os.environ.pop(key, None)
    os.environ["NO_PROXY"] = "*"
    os.environ["no_proxy"] = "*"

    doc = Document(str(docx_path))
    translator = GoogleTranslator(source=source_lang, target=target_lang)

    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)

    for paragraph in paragraphs:
        if not paragraph.text.strip():
            continue
        for run in paragraph.runs:
            text = run.text
            if not text or not text.strip():
                continue
            lead_spaces = len(text) - len(text.lstrip())
            trail_spaces = len(text) - len(text.rstrip())
            lead = text[:lead_spaces] if lead_spaces else ""
            trail = text[len(text) - trail_spaces :] if trail_spaces else ""
            core = text.strip()
            if not core:
                continue
            try:
                translated = translator.translate(core)
                if translated:
                    run.text = lead + translated + trail
            except Exception:
                continue

    doc.save(str(docx_path))


def template_error(message: str, code: int = 400):
    tool_slug = (request.form.get("_tool") or request.args.get("tool") or "").strip()
    if tool_slug in TOOL_MAP:
        return (
            render_template(
                "tool_detail.html",
                tool_slug=tool_slug,
                tool_title=TOOL_MAP[tool_slug],
                tool_desc=TOOL_DESC.get(tool_slug, ""),
                tools=TOOLS,
                category_items=CATEGORY_ITEMS,
                tool_map=TOOL_MAP,
                error=message,
                tesseract_available=TESSERACT_AVAILABLE,
            ),
            code,
        )
    return render_template("tools_list.html", tools=TOOLS, tool_desc=TOOL_DESC, category_items=CATEGORY_ITEMS, tool_map=TOOL_MAP, error=message), code


@app.get("/")
def home():
    return render_template("home.html", tools=TOOLS, tool_desc=TOOL_DESC, category_items=CATEGORY_ITEMS, tool_map=TOOL_MAP)


@app.get("/converter")
def converter_redirect():
    return redirect(url_for("tools_page"))


@app.get("/ferramentas")
def tools_page():
    return render_template("tools_list.html", tools=TOOLS, tool_desc=TOOL_DESC, category_items=CATEGORY_ITEMS, tool_map=TOOL_MAP)


@app.get("/privacidade")
def privacy_page():
    return render_template("privacy.html")


@app.get("/termos")
def terms_page():
    return render_template("terms.html")


@app.get("/contactos")
def contacts_page():
    return render_template("contacts.html")


@app.get("/sobre")
def about_page():
    return render_template("about.html")


@app.get("/ferramentas/<tool_slug>")
def tool_detail(tool_slug: str):
    if tool_slug not in TOOL_MAP:
        return template_error("Ferramenta nÃ£o encontrada.", 404)
    return render_template(
        "tool_detail.html",
        tool_slug=tool_slug,
        tool_title=TOOL_MAP[tool_slug],
        tool_desc=TOOL_DESC.get(tool_slug, ""),
        tools=TOOLS,
        category_items=CATEGORY_ITEMS,
        tool_map=TOOL_MAP,
        tesseract_available=TESSERACT_AVAILABLE,
    )


@app.post("/tools/pdf-to-word")
def pdf_to_word():
    uploaded = request.files.get("pdf_file")
    if not uploaded or not uploaded.filename.lower().endswith(".pdf"):
        return template_error("Envie um ficheiro PDF vÃ¡lido.")

    stem = Path(uploaded.filename).stem
    input_path = UPLOAD_DIR / unique_name(stem, ".pdf")
    output_path = OUTPUT_DIR / unique_name(stem, ".docx")

    translate = request.form.get("translate") == "on"
    source_lang = (request.form.get("source_lang") or "auto").strip() or "auto"
    target_lang = (request.form.get("target_lang") or "pt").strip() or "pt"

    uploaded.save(input_path)

    try:
        cv = Converter(str(input_path))
        cv.convert(str(output_path))
        cv.close()

        if translate:
            translate_docx(output_path, source_lang, target_lang)

        queue_cleanup([input_path, output_path])
        suffix = f"_{target_lang.upper()}" if translate else ""
        return send_file(str(output_path), as_attachment=True, download_name=f"{secure_filename(stem)}{suffix}.docx")
    except Exception as exc:
        return template_error(f"Falha na conversÃ£o PDF para Word: {exc}", 500)


@app.post("/tools/word-to-pdf")
def word_to_pdf():
    uploaded = request.files.get("word_file")
    if not uploaded or not uploaded.filename.lower().endswith(".docx"):
        return template_error("Envie um ficheiro DOCX vÃ¡lido.")

    stem = Path(uploaded.filename).stem
    input_path = UPLOAD_DIR / unique_name(stem, ".docx")
    output_path = OUTPUT_DIR / unique_name(stem, ".pdf")
    uploaded.save(input_path)

    try:
        converted = False

        # Windows/macOS path via docx2pdf
        if docx2pdf_convert is not None:
            try:
                docx2pdf_convert(str(input_path), str(output_path))
                converted = output_path.exists()
            except Exception:
                converted = False

        # Linux/server fallback via LibreOffice headless
        if not converted and LIBREOFFICE_PATH:
            cmd = [
                LIBREOFFICE_PATH,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_path.parent),
                str(input_path),
            ]
            proc = subprocess.run(cmd, capture_output=True, text=True, timeout=180)
            if proc.returncode == 0:
                lo_pdf = output_path.parent / f"{input_path.stem}.pdf"
                if lo_pdf.exists():
                    if lo_pdf != output_path:
                        if output_path.exists():
                            output_path.unlink(missing_ok=True)
                        lo_pdf.replace(output_path)
                    converted = True

        if not converted or not output_path.exists():
            raise RuntimeError(
                "DOCX->PDF conversion unavailable. On Linux, install LibreOffice."
            )
        queue_cleanup([input_path, output_path])
        return send_file(str(output_path), as_attachment=True, download_name=f"{secure_filename(stem)}.pdf")
    except Exception as exc:
        msg = "Word to PDF failed. Install LibreOffice on Linux or Microsoft Word/docx2pdf on Windows."
        return template_error(f"{msg} Erro: {exc}", 500)


@app.post("/tools/pdf-to-images")
def pdf_to_images():
    uploaded = request.files.get("pdf_file_images")
    if not uploaded or not uploaded.filename.lower().endswith(".pdf"):
        return template_error("Envie um PDF vÃ¡lido para converter em imagens.")

    stem = Path(uploaded.filename).stem
    input_path = UPLOAD_DIR / unique_name(stem, ".pdf")
    uploaded.save(input_path)

    scale = float(request.form.get("scale") or 2.0)
    scale = max(1.0, min(4.0, scale))

    work_dir = TEMP_DIR / unique_name(stem, "")
    work_dir.mkdir(parents=True, exist_ok=True)
    zip_path = OUTPUT_DIR / unique_name(stem, ".zip")

    try:
        doc = fitz.open(str(input_path))
        for idx, page in enumerate(doc, start=1):
            pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
            img_path = work_dir / f"{secure_filename(stem)}_p{idx}.png"
            pix.save(str(img_path))
        doc.close()

        with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for img_file in sorted(work_dir.glob("*.png")):
                zf.write(img_file, img_file.name)

        queue_cleanup([input_path, zip_path], [work_dir])
        return send_file(str(zip_path), as_attachment=True, download_name=f"{secure_filename(stem)}_images.zip")
    except Exception as exc:
        return template_error(f"Falha em PDF para imagens: {exc}", 500)


@app.post("/tools/images-to-pdf")
def images_to_pdf():
    files = request.files.getlist("image_files")
    images = [f for f in files if f and f.filename]
    if not images:
        return template_error("Envie uma ou mais imagens (PNG/JPG).")

    saved_paths: list[Path] = []
    output_path = OUTPUT_DIR / unique_name("images_to_pdf", ".pdf")

    try:
        valid_images = []
        for item in images:
            ext = Path(item.filename).suffix.lower()
            if ext not in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"}:
                continue
            img_path = UPLOAD_DIR / unique_name(Path(item.filename).stem, ext)
            item.save(img_path)
            saved_paths.append(img_path)
            valid_images.append(str(img_path))

        if not valid_images:
            return template_error("Nenhuma imagem vÃ¡lida foi enviada.")

        pdf_bytes = img2pdf.convert(valid_images)
        with output_path.open("wb") as f:
            f.write(pdf_bytes)

        queue_cleanup(saved_paths + [output_path])
        return send_file(str(output_path), as_attachment=True, download_name="images_convertido.pdf")
    except Exception as exc:
        return template_error(f"Falha em imagens para PDF: {exc}", 500)


@app.post("/tools/merge-pdf")
def merge_pdf():
    files = request.files.getlist("merge_files")
    pdfs = [f for f in files if f and f.filename and f.filename.lower().endswith(".pdf")]
    if len(pdfs) < 2:
        return template_error("Envie pelo menos 2 PDFs para juntar.")

    saved_paths: list[Path] = []
    output_path = OUTPUT_DIR / unique_name("pdf_unido", ".pdf")

    try:
        combined_doc = fitz.open()
        for item in pdfs:
            in_path = UPLOAD_DIR / unique_name(Path(item.filename).stem, ".pdf")
            item.save(in_path)
            saved_paths.append(in_path)

            with fitz.open(str(in_path)) as src:
                combined_doc.insert_pdf(src)

        combined_doc.save(str(output_path), garbage=4, deflate=True)
        combined_doc.close()

        queue_cleanup(saved_paths + [output_path])
        return send_file(str(output_path), as_attachment=True, download_name="pdf_unido.pdf")
    except Exception as exc:
        return template_error(f"Falha ao juntar PDFs: {exc}", 500)


@app.post("/tools/split-pdf")
def split_pdf():
    uploaded = request.files.get("split_file")
    if not uploaded or not uploaded.filename.lower().endswith(".pdf"):
        return template_error("Envie um PDF vÃ¡lido para dividir.")

    mode = request.form.get("split_mode", "all")
    stem = Path(uploaded.filename).stem
    input_path = UPLOAD_DIR / unique_name(stem, ".pdf")
    uploaded.save(input_path)

    try:
        doc = fitz.open(str(input_path))
        total = len(doc)

        if mode == "range":
            start = int(request.form.get("start_page") or 1)
            end = int(request.form.get("end_page") or total)
            start = max(1, min(start, total))
            end = max(start, min(end, total))

            new_doc = fitz.open()
            new_doc.insert_pdf(doc, from_page=start - 1, to_page=end - 1)
            output_path = OUTPUT_DIR / unique_name(f"{stem}_{start}_{end}", ".pdf")
            new_doc.save(str(output_path), garbage=4, deflate=True)
            new_doc.close()
            doc.close()

            queue_cleanup([input_path, output_path])
            return send_file(str(output_path), as_attachment=True, download_name=f"{secure_filename(stem)}_{start}_{end}.pdf")

        work_dir = TEMP_DIR / unique_name(stem, "")
        work_dir.mkdir(parents=True, exist_ok=True)
        zip_path = OUTPUT_DIR / unique_name(f"{stem}_split", ".zip")

        with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for i in range(total):
                page_doc = fitz.open()
                page_doc.insert_pdf(doc, from_page=i, to_page=i)
                part_name = f"{secure_filename(stem)}_p{i + 1}.pdf"
                part_path = work_dir / part_name
                page_doc.save(str(part_path), garbage=4, deflate=True)
                page_doc.close()
                zf.write(part_path, part_name)

        doc.close()
        queue_cleanup([input_path, zip_path], [work_dir])
        return send_file(str(zip_path), as_attachment=True, download_name=f"{secure_filename(stem)}_split.zip")
    except Exception as exc:
        return template_error(f"Falha ao dividir PDF: {exc}", 500)


@app.post("/tools/compress-pdf")
def compress_pdf():
    uploaded = request.files.get("compress_file")
    if not uploaded or not uploaded.filename.lower().endswith(".pdf"):
        return template_error("Envie um PDF vÃ¡lido para compressÃ£o.")

    stem = Path(uploaded.filename).stem
    input_path = UPLOAD_DIR / unique_name(stem, ".pdf")
    output_path = OUTPUT_DIR / unique_name(f"{stem}_compressed", ".pdf")
    uploaded.save(input_path)

    try:
        doc = fitz.open(str(input_path))
        doc.save(str(output_path), garbage=4, deflate=True, clean=True)
        doc.close()

        queue_cleanup([input_path, output_path])
        return send_file(str(output_path), as_attachment=True, download_name=f"{secure_filename(stem)}_compressed.pdf")
    except Exception as exc:
        return template_error(f"Falha ao comprimir PDF: {exc}", 500)


@app.post("/tools/protect-pdf")
def protect_pdf():
    uploaded = request.files.get("protect_file")
    password = (request.form.get("password") or "").strip()
    owner_password = (request.form.get("owner_password") or "").strip() or password

    if not uploaded or not uploaded.filename.lower().endswith(".pdf"):
        return template_error("Envie um PDF vÃ¡lido para proteger.")
    if not password:
        return template_error("Defina a palavra-passe do PDF.")

    stem = Path(uploaded.filename).stem
    input_path = UPLOAD_DIR / unique_name(stem, ".pdf")
    output_path = OUTPUT_DIR / unique_name(f"{stem}_protected", ".pdf")
    uploaded.save(input_path)

    try:
        doc = fitz.open(str(input_path))
        doc.save(
            str(output_path),
            encryption=fitz.PDF_ENCRYPT_AES_256,
            owner_pw=owner_password,
            user_pw=password,
            permissions=fitz.PDF_PERM_ACCESSIBILITY | fitz.PDF_PERM_PRINT | fitz.PDF_PERM_COPY,
            garbage=4,
            deflate=True,
        )
        doc.close()

        queue_cleanup([input_path, output_path])
        return send_file(str(output_path), as_attachment=True, download_name=f"{secure_filename(stem)}_protected.pdf")
    except Exception as exc:
        return template_error(f"Falha ao proteger PDF: {exc}", 500)


@app.post("/tools/unlock-pdf")
def unlock_pdf():
    uploaded = request.files.get("unlock_file")
    password = (request.form.get("unlock_password") or "").strip()

    if not uploaded or not uploaded.filename.lower().endswith(".pdf"):
        return template_error("Envie um PDF vÃ¡lido para desbloquear.")
    if not password:
        return template_error("Indique a palavra-passe para desbloquear.")

    stem = Path(uploaded.filename).stem
    input_path = UPLOAD_DIR / unique_name(stem, ".pdf")
    output_path = OUTPUT_DIR / unique_name(f"{stem}_unlocked", ".pdf")
    uploaded.save(input_path)

    try:
        doc = fitz.open(str(input_path))
        if doc.is_encrypted:
            if not doc.authenticate(password):
                doc.close()
                return template_error("Palavra-passe invÃ¡lida para este PDF.", 403)

        doc.save(str(output_path), garbage=4, deflate=True)
        doc.close()

        queue_cleanup([input_path, output_path])
        return send_file(str(output_path), as_attachment=True, download_name=f"{secure_filename(stem)}_unlocked.pdf")
    except Exception as exc:
        return template_error(f"Falha ao desbloquear PDF: {exc}", 500)


@app.post("/tools/sign-pdf")
def sign_pdf():
    uploaded = request.files.get("sign_file")
    signature_image = request.files.get("signature_image")
    signature_text = (request.form.get("signature_text") or "Assinado digitalmente").strip()

    if not uploaded or not uploaded.filename.lower().endswith(".pdf"):
        return template_error("Envie um PDF vÃ¡lido para assinar.")

    page_choice = (request.form.get("signature_page") or "last").strip()

    stem = Path(uploaded.filename).stem
    input_path = UPLOAD_DIR / unique_name(stem, ".pdf")
    output_path = OUTPUT_DIR / unique_name(f"{stem}_signed", ".pdf")
    uploaded.save(input_path)

    sign_img_path = None
    if signature_image and signature_image.filename:
        img_ext = Path(signature_image.filename).suffix.lower() or ".png"
        sign_img_path = UPLOAD_DIR / unique_name("signature", img_ext)
        signature_image.save(sign_img_path)

    try:
        doc = fitz.open(str(input_path))
        page_idx = len(doc) - 1 if page_choice == "last" else 0
        page = doc[page_idx]

        rect = page.rect
        sign_rect = fitz.Rect(rect.width - 210, rect.height - 100, rect.width - 20, rect.height - 20)
        page.draw_rect(sign_rect, color=(0.2, 0.2, 0.2), fill=(1, 1, 1), width=0.8)

        if sign_img_path and sign_img_path.exists():
            img_rect = fitz.Rect(sign_rect.x0 + 6, sign_rect.y0 + 6, sign_rect.x0 + 80, sign_rect.y1 - 6)
            page.insert_image(img_rect, filename=str(sign_img_path), keep_proportion=True)
            text_rect = fitz.Rect(sign_rect.x0 + 86, sign_rect.y0 + 8, sign_rect.x1 - 6, sign_rect.y1 - 6)
        else:
            text_rect = fitz.Rect(sign_rect.x0 + 8, sign_rect.y0 + 8, sign_rect.x1 - 8, sign_rect.y1 - 8)

        page.insert_textbox(
            text_rect,
            signature_text,
            fontsize=9,
            fontname="helv",
            color=(0.05, 0.2, 0.35),
            align=0,
        )

        doc.save(str(output_path), garbage=3, deflate=True)
        doc.close()

        cleanup_paths = [input_path, output_path]
        if sign_img_path:
            cleanup_paths.append(sign_img_path)

        queue_cleanup(cleanup_paths)
        return send_file(str(output_path), as_attachment=True, download_name=f"{secure_filename(stem)}_signed.pdf")
    except Exception as exc:
        return template_error(f"Falha ao assinar PDF: {exc}", 500)


@app.post("/tools/remove-watermark")
def remove_watermark():
    uploaded = request.files.get("watermark_file")
    watermark_text = (request.form.get("watermark_text") or "").strip()

    if not uploaded or not uploaded.filename.lower().endswith(".pdf"):
        return template_error("Envie um PDF valido para remover marca de agua.")
    if not watermark_text:
        return template_error("Indique o texto da marca de agua a remover.")

    stem = Path(uploaded.filename).stem
    input_path = UPLOAD_DIR / unique_name(stem, ".pdf")
    output_path = OUTPUT_DIR / unique_name(f"{stem}_no_watermark", ".pdf")
    uploaded.save(input_path)

    try:
        doc = fitz.open(str(input_path))
        found = 0
        for page in doc:
            areas = page.search_for(watermark_text)
            for rect in areas:
                page.add_redact_annot(rect, fill=(1, 1, 1))
                found += 1
            if areas:
                page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)

        if found == 0:
            doc.close()
            return template_error("Nao encontrei essa marca de agua no documento.")

        doc.save(str(output_path), garbage=4, deflate=True, clean=True)
        doc.close()

        queue_cleanup([input_path, output_path])
        return send_file(str(output_path), as_attachment=True, download_name=f"{secure_filename(stem)}_no_watermark.pdf")
    except Exception as exc:
        return template_error(f"Falha ao remover marca de agua: {exc}", 500)
@app.post("/tools/remove-password")
def remove_password_document():
    uploaded = request.files.get("remove_pass_file")
    password = (request.form.get("remove_pass_value") or "").strip()

    if not uploaded or not uploaded.filename.lower().endswith(".pdf"):
        return template_error("Envie um PDF vÃ¡lido para remover password.")

    stem = Path(uploaded.filename).stem
    input_path = UPLOAD_DIR / unique_name(stem, ".pdf")
    output_path = OUTPUT_DIR / unique_name(f"{stem}_no_password", ".pdf")
    uploaded.save(input_path)

    try:
        doc = fitz.open(str(input_path))
        if doc.is_encrypted:
            if not password:
                doc.close()
                return template_error("Este PDF estÃ¡ protegido. Indique a password atual.")
            if not doc.authenticate(password):
                doc.close()
                return template_error("Password invÃ¡lida para este PDF.", 403)

        doc.save(str(output_path), garbage=4, deflate=True)
        doc.close()

        queue_cleanup([input_path, output_path])
        return send_file(str(output_path), as_attachment=True, download_name=f"{secure_filename(stem)}_no_password.pdf")
    except Exception as exc:
        return template_error(f"Falha ao remover password: {exc}", 500)


@app.post("/tools/ocr")
def ocr_extract():
    uploaded = request.files.get("ocr_file")
    if not uploaded or not uploaded.filename:
        return template_error("Envie um PDF ou imagem para OCR.")

    filename = uploaded.filename.lower()
    stem = Path(uploaded.filename).stem
    ext = Path(uploaded.filename).suffix.lower()
    input_path = UPLOAD_DIR / unique_name(stem, ext)
    output_path = OUTPUT_DIR / unique_name(f"{stem}_ocr", ".txt")
    uploaded.save(input_path)

    language = (request.form.get("ocr_lang") or "por").strip() or "por"

    try:
        extracted_parts: list[str] = []

        if ext == ".pdf":
            doc = fitz.open(str(input_path))
            for i, page in enumerate(doc, start=1):
                text = page.get_text("text")
                if text and text.strip():
                    extracted_parts.append(f"\n--- PÃ¡gina {i} ---\n{text}")
                elif TESSERACT_AVAILABLE:
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                    img_path = TEMP_DIR / unique_name(f"ocr_page_{i}", ".png")
                    pix.save(str(img_path))
                    with Image.open(img_path) as img:
                        ocr_text = pytesseract.image_to_string(img, lang=language)
                    extracted_parts.append(f"\n--- PÃ¡gina {i} (OCR) ---\n{ocr_text}")
                    img_path.unlink(missing_ok=True)
            doc.close()
        elif ext in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"}:
            if not TESSERACT_AVAILABLE:
                return template_error("OCR para imagem requer Tesseract instalado no servidor.", 500)
            with Image.open(input_path) as img:
                extracted_parts.append(pytesseract.image_to_string(img, lang=language))
        else:
            return template_error("Formato invÃ¡lido para OCR. Use PDF ou imagem.")

        text_out = "\n".join(extracted_parts).strip()
        if not text_out:
            text_out = "Nenhum texto detetado no ficheiro."

        output_path.write_text(text_out, encoding="utf-8")

        queue_cleanup([input_path, output_path])
        return send_file(str(output_path), as_attachment=True, download_name=f"{secure_filename(stem)}_ocr.txt")
    except Exception as exc:
        return template_error(f"Falha no OCR/extraÃ§Ã£o de texto: {exc}", 500)


@app.post("/tools/transcribe-audio")
def transcribe_audio():
    uploaded = request.files.get("audio_file")
    lang = (request.form.get("audio_lang") or "pt-PT").strip() or "pt-PT"

    if not uploaded or not uploaded.filename:
        return template_error("Envie um ficheiro de audio ou video para transcrever.")

    ext = Path(uploaded.filename).suffix.lower()
    if ext not in {".mp3", ".wav", ".m4a", ".ogg", ".flac", ".mp4", ".webm", ".mov"}:
        return template_error("Formato nao suportado para transcricao.")
    if AudioSegment is None or sr is None:
        return template_error("Transcricao indisponivel: instale pydub e SpeechRecognition no servidor.", 500)

    stem = Path(uploaded.filename).stem
    input_path = RUNTIME_DIR / unique_name(stem, ext)
    wav_dir = RUNTIME_DIR / unique_name("transcribe_chunks", "")
    output_path = OUTPUT_DIR / unique_name(f"{stem}_transcricao", ".txt")
    wav_dir.mkdir(parents=True, exist_ok=True)
    uploaded.save(input_path)

    try:
        for key in ("http_proxy", "https_proxy", "HTTP_PROXY", "HTTPS_PROXY", "ALL_PROXY", "all_proxy"):
            os.environ.pop(key, None)
        os.environ["NO_PROXY"] = "*"
        os.environ["no_proxy"] = "*"

        audio = AudioSegment.from_file(input_path)
        recognizer = sr.Recognizer()
        chunk_ms = 45000
        texts: list[str] = []

        idx = 1
        for start in range(0, len(audio), chunk_ms):
            chunk = audio[start : start + chunk_ms]
            wav_path = wav_dir / f"chunk_{idx}.wav"
            chunk.export(wav_path, format="wav")

            with sr.AudioFile(str(wav_path)) as source:
                data = recognizer.record(source)
            try:
                text = recognizer.recognize_google(data, language=lang)
            except sr.UnknownValueError:
                text = ""
            except sr.RequestError as req_err:
                return template_error(f"Falha no servico de transcricao: {req_err}", 502)

            texts.append(text)
            idx += 1

        transcript = "\n".join(part for part in texts if part.strip()).strip()
        if not transcript:
            transcript = "Nao foi possivel transcrever audio reconhecivel."

        output_path.write_text(transcript, encoding="utf-8")
        queue_cleanup([input_path, output_path], [wav_dir])
        return send_file(str(output_path), as_attachment=True, download_name=f"{secure_filename(stem)}_transcricao.txt")
    except Exception as exc:
        return template_error(f"Falha ao transcrever audio: {exc}", 500)


@app.post("/tools/translate")
def translate_document():
    uploaded = request.files.get("translate_file")
    if not uploaded or not uploaded.filename:
        return template_error("Envie um ficheiro para traduzir (DOCX, PDF ou TXT).")

    source_lang = (request.form.get("tr_source") or "auto").strip() or "auto"
    target_lang = (request.form.get("tr_target") or "pt").strip() or "pt"

    stem = Path(uploaded.filename).stem
    ext = Path(uploaded.filename).suffix.lower()
    input_path = UPLOAD_DIR / unique_name(stem, ext)
    uploaded.save(input_path)

    try:
        for key in ("http_proxy", "https_proxy", "HTTP_PROXY", "HTTPS_PROXY", "ALL_PROXY", "all_proxy"):
            os.environ.pop(key, None)
        os.environ["NO_PROXY"] = "*"
        os.environ["no_proxy"] = "*"

        if ext == ".docx":
            output_path = OUTPUT_DIR / unique_name(f"{stem}_{target_lang}", ".docx")
            shutil.copy2(input_path, output_path)
            translate_docx(output_path, source_lang, target_lang)
            queue_cleanup([input_path, output_path])
            return send_file(str(output_path), as_attachment=True, download_name=f"{secure_filename(stem)}_{target_lang}.docx")

        if ext == ".pdf":
            docx_path = OUTPUT_DIR / unique_name(f"{stem}_translated", ".docx")
            cv = Converter(str(input_path))
            cv.convert(str(docx_path))
            cv.close()
            translate_docx(docx_path, source_lang, target_lang)
            queue_cleanup([input_path, docx_path])
            return send_file(str(docx_path), as_attachment=True, download_name=f"{secure_filename(stem)}_{target_lang}.docx")

        if ext == ".txt":
            output_path = OUTPUT_DIR / unique_name(f"{stem}_{target_lang}", ".txt")
            translator = GoogleTranslator(source=source_lang, target=target_lang)
            source_text = input_path.read_text(encoding="utf-8", errors="ignore")
            translated = translator.translate(source_text)
            output_path.write_text(translated or "", encoding="utf-8")
            queue_cleanup([input_path, output_path])
            return send_file(str(output_path), as_attachment=True, download_name=f"{secure_filename(stem)}_{target_lang}.txt")

        return template_error("Formato nÃ£o suportado. Use DOCX, PDF ou TXT.")
    except Exception as exc:
        return template_error(f"Falha na traduÃ§Ã£o de documento: {exc}", 500)


@app.post("/tools/youtube-to-mp3")
def youtube_to_mp3():
    url = (request.form.get("youtube_url") or "").strip()
    if not url:
        return template_error("Indique um link de YouTube.")
    if yt_dlp is None:
        return template_error("YouTube indisponivel: instale yt-dlp no servidor.", 500)
    if not FFMPEG_PATH:
        return template_error("ffmpeg nao encontrado no servidor. Necessario para MP3.", 500)

    work_dir = RUNTIME_DIR / unique_name("yt_mp3", "")
    work_dir.mkdir(parents=True, exist_ok=True)
    out_template = str(work_dir / "%(title)s [%(id)s].%(ext)s")
    expected_mp3 = None
    video_title = ""

    try:
        opts = {
            "format": "bestaudio/best",
            "outtmpl": out_template,
            "proxy": "",
            "ffmpeg_location": FFMPEG_PATH,
            "postprocessors": [{"key": "FFmpegExtractAudio", "preferredcodec": "mp3", "preferredquality": "192"}],
            "nopart": True,
            "quiet": True,
            "no_warnings": True,
            "noplaylist": True,
        }
        with yt_dlp.YoutubeDL(opts) as ydl:
            info = ydl.extract_info(url, download=True)
            video_title = info.get("title") or ""
            source_file = Path(ydl.prepare_filename(info))
            expected_mp3 = source_file.with_suffix(".mp3")

        if not expected_mp3 or not expected_mp3.exists():
            matches = sorted(work_dir.glob("*.mp3"), key=lambda p: p.stat().st_mtime, reverse=True)
            if not matches:
                return template_error("Nao foi possivel gerar MP3 para este link.", 500)
            expected_mp3 = matches[0]

        queue_cleanup([expected_mp3], [work_dir])
        download_name = safe_download_name(video_title, ".mp3", expected_mp3.stem)
        return send_file(str(expected_mp3), as_attachment=True, download_name=download_name)
    except Exception as exc:
        return template_error(f"Falha em YouTube para MP3: {exc}", 500)


@app.post("/tools/youtube-to-mp4")
def youtube_to_mp4():
    url = (request.form.get("youtube_url_mp4") or "").strip()
    if not url:
        return template_error("Indique um link de YouTube.")
    if yt_dlp is None:
        return template_error("YouTube indisponivel: instale yt-dlp no servidor.", 500)

    work_dir = RUNTIME_DIR / unique_name("yt_mp4", "")
    work_dir.mkdir(parents=True, exist_ok=True)
    out_template = str(work_dir / "%(title)s [%(id)s].%(ext)s")
    expected_mp4 = None
    video_title = ""

    try:
        opts = {
            "format": "bestvideo+bestaudio/best",
            "outtmpl": out_template,
            "merge_output_format": "mp4",
            "proxy": "",
            "ffmpeg_location": FFMPEG_PATH,
            "nopart": True,
            "quiet": True,
            "no_warnings": True,
            "noplaylist": True,
        }
        with yt_dlp.YoutubeDL(opts) as ydl:
            info = ydl.extract_info(url, download=True)
            video_title = info.get("title") or ""
            source_file = Path(ydl.prepare_filename(info))
            expected_mp4 = source_file.with_suffix(".mp4")

        if not expected_mp4 or not expected_mp4.exists():
            matches = sorted(work_dir.glob("*.mp4"), key=lambda p: p.stat().st_mtime, reverse=True)
            if not matches:
                return template_error("Nao foi possivel gerar MP4 para este link.", 500)
            expected_mp4 = matches[0]

        queue_cleanup([expected_mp4], [work_dir])
        download_name = safe_download_name(video_title, ".mp4", expected_mp4.stem)
        return send_file(str(expected_mp4), as_attachment=True, download_name=download_name)
    except Exception as exc:
        return template_error(f"Falha em YouTube para MP4: {exc}", 500)


@app.post("/tools/youtube-playlist-to-mp3")
def youtube_playlist_to_mp3():
    url = (request.form.get("youtube_playlist_url_mp3") or "").strip()
    if not url:
        return template_error("Indique um link de playlist do YouTube.")
    if yt_dlp is None:
        return template_error("YouTube indisponivel: instale yt-dlp no servidor.", 500)
    if not FFMPEG_PATH:
        return template_error("ffmpeg nao encontrado no servidor. Necessario para MP3.", 500)

    work_dir = RUNTIME_DIR / unique_name("yt_playlist_mp3", "")
    work_dir.mkdir(parents=True, exist_ok=True)
    out_template = str(work_dir / "%(playlist_index)03d - %(title)s [%(id)s].%(ext)s")
    playlist_title = "youtube_playlist_mp3"
    zip_path = OUTPUT_DIR / unique_name("youtube_playlist_mp3", ".zip")

    try:
        opts = {
            "format": "bestaudio/best",
            "outtmpl": out_template,
            "proxy": "",
            "ffmpeg_location": FFMPEG_PATH,
            "postprocessors": [{"key": "FFmpegExtractAudio", "preferredcodec": "mp3", "preferredquality": "192"}],
            "nopart": True,
            "quiet": True,
            "no_warnings": True,
            "noplaylist": False,
        }
        with yt_dlp.YoutubeDL(opts) as ydl:
            info = ydl.extract_info(url, download=True)
            playlist_title = info.get("title") or playlist_title

        files = sorted(work_dir.glob("*.mp3"))
        if not files:
            return template_error("Nao foi possivel gerar MP3 para esta playlist.", 500)

        zip_files(files, zip_path)
        queue_cleanup([zip_path], [work_dir])
        return send_file(
            str(zip_path),
            as_attachment=True,
            download_name=safe_download_name(playlist_title, ".zip", "youtube_playlist_mp3"),
        )
    except Exception as exc:
        return template_error(f"Falha em playlist YouTube para MP3: {exc}", 500)


@app.post("/tools/youtube-playlist-to-mp4")
def youtube_playlist_to_mp4():
    url = (request.form.get("youtube_playlist_url_mp4") or "").strip()
    if not url:
        return template_error("Indique um link de playlist do YouTube.")
    if yt_dlp is None:
        return template_error("YouTube indisponivel: instale yt-dlp no servidor.", 500)

    work_dir = RUNTIME_DIR / unique_name("yt_playlist_mp4", "")
    work_dir.mkdir(parents=True, exist_ok=True)
    out_template = str(work_dir / "%(playlist_index)03d - %(title)s [%(id)s].%(ext)s")
    playlist_title = "youtube_playlist_mp4"
    zip_path = OUTPUT_DIR / unique_name("youtube_playlist_mp4", ".zip")

    try:
        opts = {
            "format": "bestvideo+bestaudio/best",
            "outtmpl": out_template,
            "merge_output_format": "mp4",
            "proxy": "",
            "ffmpeg_location": FFMPEG_PATH,
            "nopart": True,
            "quiet": True,
            "no_warnings": True,
            "noplaylist": False,
        }
        with yt_dlp.YoutubeDL(opts) as ydl:
            info = ydl.extract_info(url, download=True)
            playlist_title = info.get("title") or playlist_title

        files = sorted(work_dir.glob("*.mp4"))
        if not files:
            return template_error("Nao foi possivel gerar MP4 para esta playlist.", 500)

        zip_files(files, zip_path)
        queue_cleanup([zip_path], [work_dir])
        return send_file(
            str(zip_path),
            as_attachment=True,
            download_name=safe_download_name(playlist_title, ".zip", "youtube_playlist_mp4"),
        )
    except Exception as exc:
        return template_error(f"Falha em playlist YouTube para MP4: {exc}", 500)


@app.post("/tools/spotify-to-mp3")
def spotify_to_mp3():
    url = (request.form.get("spotify_url") or "").strip()
    if not url:
        return template_error("Indique um link de Spotify.")

    work_dir = RUNTIME_DIR / unique_name("spotify_dl", "")
    work_dir.mkdir(parents=True, exist_ok=True)

    try:
        cmd = [
            "python",
            "-m",
            "spotdl",
            url,
            "--output",
            str(work_dir),
            "--path-template",
            "{title}.{ext}",
            "--output-format",
            "mp3",
            "--ffmpeg",
            FFMPEG_PATH or "ffmpeg",
        ]
        env = os.environ.copy()
        for key in ("http_proxy", "https_proxy", "HTTP_PROXY", "HTTPS_PROXY", "ALL_PROXY", "all_proxy"):
            env.pop(key, None)
        env["NO_PROXY"] = "*"
        env["no_proxy"] = "*"
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=600, env=env)
        if proc.returncode != 0:
            msg = (proc.stderr or proc.stdout or "Erro desconhecido no spotdl.").strip()
            return template_error(f"Falha em Spotify para MP3: {msg}", 500)

        mp3_files = sorted(work_dir.glob("*.mp3"), key=lambda p: p.stat().st_mtime, reverse=True)
        if not mp3_files:
            return template_error("Nao foi gerado MP3 para este link de Spotify.", 500)

        result_path = mp3_files[0]
        queue_cleanup([result_path], [work_dir])
        download_name = safe_download_name(result_path.stem, ".mp3", "spotify_audio")
        return send_file(str(result_path), as_attachment=True, download_name=download_name)
    except subprocess.TimeoutExpired:
        return template_error("Tempo limite excedido no download do Spotify.", 504)
    except Exception as exc:
        return template_error(f"Falha em Spotify para MP3: {exc}", 500)


@app.post("/tools/spotify-playlist-to-mp3")
def spotify_playlist_to_mp3():
    url = (request.form.get("spotify_playlist_url") or "").strip()
    if not url:
        return template_error("Indique um link de playlist do Spotify.")

    work_dir = RUNTIME_DIR / unique_name("spotify_playlist_dl", "")
    work_dir.mkdir(parents=True, exist_ok=True)
    zip_path = OUTPUT_DIR / unique_name("spotify_playlist_mp3", ".zip")

    try:
        cmd = [
            "python",
            "-m",
            "spotdl",
            url,
            "--output",
            str(work_dir),
            "--path-template",
            "{artists} - {title}.{ext}",
            "--output-format",
            "mp3",
            "--ffmpeg",
            FFMPEG_PATH or "ffmpeg",
        ]
        env = os.environ.copy()
        for key in ("http_proxy", "https_proxy", "HTTP_PROXY", "HTTPS_PROXY", "ALL_PROXY", "all_proxy"):
            env.pop(key, None)
        env["NO_PROXY"] = "*"
        env["no_proxy"] = "*"
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=1800, env=env)
        if proc.returncode != 0:
            msg = (proc.stderr or proc.stdout or "Erro desconhecido no spotdl.").strip()
            return template_error(f"Falha em playlist Spotify para MP3: {msg}", 500)

        files = sorted(work_dir.glob("*.mp3"))
        if not files:
            return template_error("Nao foi gerado MP3 para esta playlist do Spotify.", 500)

        zip_files(files, zip_path)
        queue_cleanup([zip_path], [work_dir])
        return send_file(str(zip_path), as_attachment=True, download_name="spotify_playlist_mp3.zip")
    except subprocess.TimeoutExpired:
        return template_error("Tempo limite excedido no download da playlist Spotify.", 504)
    except Exception as exc:
        return template_error(f"Falha em playlist Spotify para MP3: {exc}", 500)


@app.errorhandler(413)
def file_too_large(_):
    return template_error("Ficheiro demasiado grande. Limite atual: 40MB.", 413)


@app.get("/robots.txt")
def robots_txt():
    base_url = os.environ.get("BASE_URL", request.host_url.rstrip("/"))
    content = f"User-agent: *\nAllow: /\nSitemap: {base_url}/sitemap.xml\n"
    return content, 200, {"Content-Type": "text/plain"}


@app.get("/sitemap.xml")
def sitemap_xml():
    base_url = os.environ.get("BASE_URL", request.host_url.rstrip("/"))
    # Filter out media category for AdSense approval phase
    hidden_slugs = CATEGORY_ITEMS.get("media", [])
    
    pages = [
        {"loc": f"{base_url}/", "lastmod": "2025-03-10", "priority": "1.0"},
        {"loc": f"{base_url}/ferramentas", "lastmod": "2025-03-10", "priority": "0.9"},
    ]
    for slug, _ in TOOLS:
        if slug not in hidden_slugs:
            pages.append({"loc": f"{base_url}/ferramentas/{slug}", "lastmod": "2025-03-10", "priority": "0.8"})

    xml = '<?xml version="1.0" encoding="UTF-8"?>\n'
    xml += '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">\n'
    for page in pages:
        xml += "  <url>\n"
        xml += f'    <loc>{page["loc"]}</loc>\n'
        xml += f'    <lastmod>{page["lastmod"]}</lastmod>\n'
        xml += f'    <priority>{page["priority"]}</priority>\n'
        xml += "  </url>\n"
    xml += "</urlset>"
    return xml, 200, {"Content-Type": "application/xml"}


if __name__ == "__main__":
    port = int(os.environ.get("PORT", "5000"))
    app.run(host="0.0.0.0", port=port, debug=True)
